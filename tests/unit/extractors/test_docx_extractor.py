# Copyright (c) 2025 takotime808

import pytest
from pathlib import Path
from docx import Document
from unifile.extractors.docx_extractor import DocxExtractor

def _build_docx(path: Path):
    doc = Document()
    doc.add_paragraph("Paragraph one")
    doc.add_paragraph("Paragraph two")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "C"
    table.rows[1].cells[1].text = "D"
    doc.save(path)

def test_docx_extractor_paragraphs_and_tables(tmp_path):
    p = tmp_path / "sample.docx"
    _build_docx(p)
    ext = DocxExtractor()
    rows = ext.extract(p)
    # expect a body row and at least one table row
    types = [r.unit_type for r in rows]
    assert "file" in types or "body" in [r.unit_id for r in rows]
    assert "table" in types
    # body contains the paragraphs concatenated
    body = [r for r in rows if r.unit_type == "file"]
    assert body and "Paragraph one" in body[0].content
    # table content should include tab-separated cells
    tables = [r for r in rows if r.unit_type == "table"]
    assert tables and "A\tB" in tables[0].content
