# Copyright (c) 2025 takotime808

from __future__ import annotations

from pathlib import Path
from typing import List
from docx import Document

from unifile.extractors.base import (
    BaseExtractor,
    make_row,
    Row,
)


class DocxExtractor(BaseExtractor):
    """
    DOCX paragraph + table extractor.

    This extractor reads:
      - **Paragraph text**: concatenated into a single "file/body" unit.
      - **Tables**: each table is serialized as tab-delimited rows and emitted
        as a separate "table/<index>" unit.

    It inherits a robust :meth:`extract` implementation from :class:`BaseExtractor`
    that performs file validation and wraps unexpected exceptions in a single
    error :class:`Row`. The real work is implemented in :meth:`_extract`.

    Output Rows
    -----------
    * Paragraphs (if any):
        - file_type: "docx"
        - unit_type: "file"
        - unit_id:   "body"
        - metadata:  {"sections": "paragraphs"}
    * Each table:
        - file_type: "docx"
        - unit_type: "table"
        - unit_id:   "<table_index>"
        - content:   Tab-delimited lines (one line per table row)
        - metadata:  {"table_index": <int>}

    Notes
    -----
    - Empty paragraphs are skipped.
    - Empty tables (i.e., after trimming) are not emitted.
    """

    supported_extensions = ["docx"]

    def _extract(self, path: Path) -> List[Row]:
        """
        Parse a DOCX file into standardized rows.

        Parameters
        ----------
        path
            Path to a `.docx` file. Validation (existence/is-file) is already
            handled by :class:`BaseExtractor.extract`.

        Returns
        -------
        list[Row]
            Standardized rows for paragraphs and tables (see class docstring).
        """
        doc = Document(str(path))

        # Collect paragraph text (skip empty)
        paragraph_texts: list[str] = []
        for p in doc.paragraphs:
            if p.text:
                paragraph_texts.append(p.text)

        # Collect tables (serialize cells as tab-delimited)
        table_blobs: list[tuple[int, str]] = []
        for ti, t in enumerate(doc.tables):
            lines: list[str] = []
            for row in t.rows:
                lines.append("\t".join((cell.text or "") for cell in row.cells))
            table_text = "\n".join(lines).strip()
            if table_text:
                table_blobs.append((ti, table_text))

        rows: List[Row] = []
        if paragraph_texts:
            rows.append(
                make_row(
                    path=path,
                    file_type="docx",
                    unit_type="file",
                    unit_id="body",
                    content="\n".join(paragraph_texts).strip(),
                    metadata={"sections": "paragraphs"},
                    status="ok",
                )
            )

        for ti, blob in table_blobs:
            rows.append(
                make_row(
                    path=path,
                    file_type="docx",
                    unit_type="table",
                    unit_id=str(ti),
                    content=blob,
                    metadata={"table_index": ti},
                    status="ok",
                )
            )

        return rows
